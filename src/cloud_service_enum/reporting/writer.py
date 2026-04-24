"""Report writing for a single :class:`EnumerationRun`.

Each format has one small helper; the public entry point
:func:`write_reports` accepts a list of formats, writes each in turn,
and returns the filesystem paths produced.
"""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from datetime import datetime
from enum import StrEnum
from pathlib import Path
from typing import Any
from xml.etree.ElementTree import Element, ElementTree, SubElement

from cloud_service_enum.core.models import EnumerationRun, MultiAccountRun, ServiceResult


class ReportFormat(StrEnum):
    JSON = "json"
    CSV = "csv"
    XLSX = "xlsx"
    XML = "xml"
    DOCX = "docx"
    PDF = "pdf"


@dataclass(frozen=True)
class ReportWriter:
    """Write a single :class:`EnumerationRun` to one format."""

    run: EnumerationRun
    output_dir: Path

    def write(self, fmt: ReportFormat) -> Path:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        ts = self.run.started_at.strftime("%Y%m%d-%H%M%S")
        suffix = f"-{self.run.profile}" if self.run.profile else ""
        stem = f"{self.run.provider.value}{suffix}-{ts}"
        path = self.output_dir / f"{stem}.{fmt.value}"
        _WRITERS[fmt](self.run, path)
        return path


def write_reports(
    run: EnumerationRun | MultiAccountRun,
    output_dir: str | Path,
    formats: list[ReportFormat],
) -> list[Path]:
    """Write a run to every format in ``formats`` and return the paths.

    A :class:`MultiAccountRun` fans out into one report set per member
    run plus, for JSON only, a single combined document so downstream
    tools can consume the aggregate shape directly.
    """
    output = Path(output_dir)
    if isinstance(run, MultiAccountRun):
        paths: list[Path] = []
        for member in run.accounts:
            paths.extend(write_reports(member, output, formats))
        if ReportFormat.JSON in formats:
            paths.append(_multi_json(run, output))
        return paths
    writer = ReportWriter(run=run, output_dir=output)
    return [writer.write(fmt) for fmt in formats]


def _multi_json(multi: MultiAccountRun, output_dir: Path) -> Path:
    """Serialise a :class:`MultiAccountRun` to a single combined JSON doc."""
    output_dir.mkdir(parents=True, exist_ok=True)
    ts = multi.started_at.strftime("%Y%m%d-%H%M%S")
    path = output_dir / f"{multi.provider.value}-multi-{ts}.json"
    path.write_text(multi.model_dump_json(indent=2, exclude_none=True), encoding="utf-8")
    return path


def _json(run: EnumerationRun, path: Path) -> None:
    path.write_text(run.model_dump_json(indent=2, exclude_none=True), encoding="utf-8")


def _csv(run: EnumerationRun, path: Path) -> None:
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["service", "kind", "id", "name", "region", "payload"])
        for svc in run.services:
            for res in svc.resources:
                writer.writerow(
                    [
                        svc.service,
                        res.get("kind", ""),
                        res.get("id", ""),
                        res.get("name", ""),
                        res.get("region", ""),
                        json.dumps({k: v for k, v in res.items() if k not in {"kind", "id", "name", "region"}}, default=str),
                    ]
                )


def _xml(run: EnumerationRun, path: Path) -> None:
    root = Element("enumeration", provider=run.provider.value)
    SubElement(root, "identity").text = json.dumps(run.identity, default=str)
    services = SubElement(root, "services")
    for svc in run.services:
        s = SubElement(services, "service", name=svc.service, count=str(svc.count))
        for res in svc.resources:
            SubElement(s, "resource").text = json.dumps(res, default=str)
        for err in svc.errors:
            SubElement(s, "error").text = err
    ElementTree(root).write(path, encoding="utf-8", xml_declaration=True)


def _xlsx(run: EnumerationRun, path: Path) -> None:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise RuntimeError("install extras 'reports' for XLSX support") from exc

    wb = Workbook()
    summary = wb.active
    summary.title = "summary"
    summary.append(["provider", "service", "resources", "errors", "duration_s"])
    for svc in run.services:
        summary.append([run.provider.value, svc.service, svc.count, len(svc.errors), svc.duration_s])

    for svc in run.services:
        sheet = wb.create_sheet(svc.service[:31])
        columns = _columns_for(svc)
        sheet.append(columns)
        for res in svc.resources:
            sheet.append([_cell(res.get(c)) for c in columns])
    wb.save(path)


def _docx(run: EnumerationRun, path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("install extras 'reports' for DOCX support") from exc

    doc = Document()
    doc.add_heading(f"{run.provider.value.upper()} enumeration report", level=0)
    doc.add_paragraph(f"Started: {run.started_at.isoformat()}")
    doc.add_paragraph(f"Duration: {run.duration_s}s")
    doc.add_paragraph(f"Principal: {run.identity.get('principal', '?')}")

    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Service", "Resources", "Errors"
    for svc in run.services:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = svc.service, str(svc.count), str(len(svc.errors))

    for svc in run.services:
        doc.add_heading(svc.service, level=2)
        if not svc.resources:
            doc.add_paragraph("(no resources)", style="Intense Quote")
        for res in svc.resources[:200]:
            doc.add_paragraph(json.dumps(res, indent=2, default=str), style="Normal")
        if svc.errors:
            doc.add_paragraph("Errors: " + "; ".join(svc.errors), style="Intense Quote")
    doc.save(path)


def _pdf(run: EnumerationRun, path: Path) -> None:
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise RuntimeError("install extras 'reports' for PDF support") from exc

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, f"{run.provider.value.upper()} Enumeration Report", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, f"Started: {run.started_at.isoformat()}", ln=True)
    pdf.cell(0, 6, f"Duration: {run.duration_s}s", ln=True)
    pdf.cell(0, 6, f"Principal: {run.identity.get('principal', '?')}", ln=True)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Summary", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for svc in run.services:
        pdf.cell(
            0,
            5,
            f"- {svc.service}: {svc.count} resources, {len(svc.errors)} errors, {svc.duration_s}s",
            ln=True,
        )
    for svc in run.services:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 8, svc.service, ln=True)
        pdf.set_font("Courier", "", 8)
        for res in svc.resources[:50]:
            text = json.dumps(res, indent=2, default=str)
            pdf.multi_cell(0, 4, text)
            pdf.ln(1)
    pdf.output(str(path))


def _columns_for(svc: ServiceResult) -> list[str]:
    preferred = ["kind", "id", "name", "region", "arn"]
    seen = set(preferred)
    for res in svc.resources:
        for key in res:
            if key not in seen:
                preferred.append(key)
                seen.add(key)
    return preferred


def _cell(value: Any) -> Any:
    if isinstance(value, (dict, list)):
        return json.dumps(value, default=str)
    if isinstance(value, datetime):
        return value.isoformat()
    return value


_WRITERS = {
    ReportFormat.JSON: _json,
    ReportFormat.CSV: _csv,
    ReportFormat.XML: _xml,
    ReportFormat.XLSX: _xlsx,
    ReportFormat.DOCX: _docx,
    ReportFormat.PDF: _pdf,
}
